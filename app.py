# main.py

import yaml
import os
import glob
import pandas as pd
from docx import Document
from docx.shared import Inches
from spire.pdf import *
import pytesseract
from utils import (preprocess_image, find_contours, transform_carplate, 
                  extract_fine_amount, send_email)

def load_config():
    with open('config.yaml', 'r') as f:
        return yaml.safe_load(f)

def extract_text(pdf_path):
    """Extracts text from a Spire.PDF document."""
    doc = PdfDocument()
    doc.LoadFromFile(pdf_path)
    page0 = doc.Pages[0]
    page1 = doc.Pages[1]
    text = []
    for page in [page0,page1]:
        text_extractor = PdfTextExtractor(page)
        extract_options = PdfTextExtractOptions()
        text.append(text_extractor.ExtractText(extract_options))
    text = ' '.join(text)
    return text

def get_reg_photo(pdf_path):
    """Extracts photos from the PDF using Spire.PDF."""
    doc = PdfDocument()
    doc.LoadFromFile(pdf_path)
    page = doc.Pages.get_Item(1)
    image_helper = PdfImageHelper()
    image_infos = image_helper.GetImagesInfo(page)
    reg_number_image = image_infos[0].Image
    car_photo = image_infos[1].Image
    reg_number_filename = f'/content/numbers/number_{os.path.splitext(os.path.basename(pdf_path))[0]}.png'
    car_photo_filename = f'/content/numbers/car_{os.path.splitext(os.path.basename(pdf_path))[0]}.png'
    savedir = "/content/numbers/"
    os.makedirs(savedir, exist_ok=True)
    os.chdir(savedir)
    reg_number_image.Save(reg_number_filename)
    car_photo.Save(car_photo_filename)
    print(f'Images from {os.path.splitext(os.path.basename(pdf_path))[0]} are extracted !')
    return reg_number_filename, car_photo_filename

def get_text_info(text):
    """Extracts specific data from text using regular expressions."""
    data = {}
    data['ПОСТАНОВЛЕНИЕ'] = re.search(r'ПОСТАНОВЛЕНИЕ (\d+)', text).group(1) if re.search(r'ПОСТАНОВЛЕНИЕ (\d+)', text) else None
    data['дата_нарушения'] = re.search(r'УСТАНОВИЛ:\n.*(\d{2}\.\d{2}\.\d{4})', text).group(1) if re.search(r'УСТАНОВИЛ:\n.*(\d{2}\.\d{2}\.\d{4})', text) else None
    data['время_нарушения'] = re.search(r'(\d{2}:\d{2}:\d{2})', text).group(1) if re.search(r'(\d{2}:\d{2}:\d{2})', text) else None
    data['адрес'] = re.sub(r'\n\s+', ' ', re.search(r'по адресу(.*?)водитель', text, re.DOTALL).group(1)).strip() if re.search(r'по адресу(.*?)водитель', text, re.DOTALL) else None
    data['номер_тс'] = re.search(r'государственный регистрационный знак (\w+)', text).group(1) if re.search(r'государственный регистрационный знак (\w+)', text) else None
    data['сумма_штрафа'] = re.search(r'штрафа в размере (\b(\d+(?:[\xa0 ]\d{3})*)\b) руб', text).group(1).replace('\xa0', '')
    data['номер_свидетельства'] = re.search(r'свидетельством о регистрации ТС №(\d+)', text).group(1) if re.search(r'свидетельством о регистрации ТС №(\d+)', text) else None
    data['ИГР'] = re.search(r'Идентификация государственного регистрационного знака:\s*(([А-Я]{1}\d{3}[А-Я]{2}\d{2,3})|([А-Я]{2}\d{4}\d{2,3}))', text).group(1) if re.search(r'Идентификация государственного регистрационного знака:\s*(([А-Я]{1}\d{3}[А-Я]{2}\d{2,3})|([А-Я]{2}\d{4}\d{2,3}))', text) else None
    return data

def tesseract_recognizer(image, psm=7):
    """OCR of the car plate number with tesseract."""
    custom_config = (f'--oem 3 --psm {psm} '
                    r'-c tessedit_char_whitelist=АВЕКМНОРСТУХ0123456789')
    ocr_result = pytesseract.image_to_string(image, lang='rus', config=custom_config)
    ocr_result_cleaned = ocr_result.strip()
    ocr_list = list(ocr_result_cleaned)
    corrected_text = ''.join(ocr_list)
    print("Recognized Plate Number:", corrected_text)
    return corrected_text

def process_pdf(pdf_path):
    """Processes the PDF and extracts information."""
    # Extract text and images
    text = extract_text(pdf_path)
    images = get_reg_photo(pdf_path)
    acc_info = get_text_info(text)

    # Preprocess and recognize car plate number
    dilated = preprocess_image(images[0])
    cont = find_contours(dilated)
    transformed_carplate = transform_carplate(dilated, cont)
    recognized_number = tesseract_recognizer(transformed_carplate)

    # Create the dictionary
    data_dict = {
        'filename': os.path.splitext(os.path.basename(pdf_path))[0],
        'car_photo_filename': images[1],
        'regnumber_filename': images[0],
        'ocr_result': recognized_number,
        **acc_info
    }
    return data_dict

def create_table(data_dicts):
    """Creates a pandas DataFrame from a list of data dictionaries."""
    df = pd.DataFrame(data_dicts)
    df['Row number'] = df.index + 1
    df['сумма_штрафа_greater_5000'] = df['сумма_штрафа'].astype(int) > 5000
    df['ocr_match'] = df['ИГР'] == df['ocr_result']

    # Embed images using HTML tags
    df['car_photo'] = df['car_photo_filename'].apply(lambda x: f'<img src="{x}" width="100">')
    df['regnumber_photo'] = df['regnumber_filename'].apply(lambda x: f'<img src="{x}" width="100">')

    # Reorder columns
    columns_order = ['Row number', 'filename', 'ПОСТАНОВЛЕНИЕ', 'дата_нарушения', 
                    'время_нарушения', 'адрес', 'номер_тс', 'сумма_штрафа', 
                    'сумма_штрафа_greater_5000', 'номер_свидетельства',
                    'car_photo_filename', 'regnumber_filename', 'ocr_match', 'ИГР']
    df = df[columns_order]
    return df

def export_to_docx(df, filename="output_table.docx"):
    """Exports the DataFrame to a docx file with images."""
    document = Document()
    table = document.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            if i in [df.columns.get_loc('car_photo_filename'), df.columns.get_loc('regnumber_filename')]:
                image_path = row[df.columns[i]]
                try:
                    run = row_cells[i].paragraphs[0].add_run()
                    run.add_picture(image_path, width=Inches(1.5))
                except:
                    row_cells[i].text = "Error loading image"
            else:
                row_cells[i].text = str(value)

    document.save(filename)

def main():
    config = load_config()
    
    # Process PDFs from the specified directory
    pdf_files = glob.glob(os.path.join(config['google_drive']['download_dir'], '*.pdf'))
    acc_dicts = []
    
    for pdf_file in pdf_files:
        acc_dict = process_pdf(pdf_file)
        acc_dicts.append(acc_dict)
    
    # Create and export results
    df = create_table(acc_dicts)
    export_to_docx(df)
    
    # Send notifications based on conditions
    for data in acc_dicts:
        if data['ИГР'] == data['ocr_result'] and int(data['сумма_штрафа']) > 5000:
            send_email(config['email'], data, config['email']['recipients']['cargo'])
        else:
            send_email(config['email'], data, config['email']['recipients']['tech'])

if __name__ == "__main__":
    main()


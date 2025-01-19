# src/processors.py

import cv2
import numpy as np
import pandas as pd
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from docx import Document
from docx.shared import Inches
from spire.pdf.common import *
from spire.pdf import *

from .exceptions import (
    PDFProcessingError,
    OCRError,
    ImageProcessingError,
    EmailError,
    ConfigurationError
)
from .ocr.recognizer import OCRRecognizer

class PDFProcessor:
    """Handles PDF document processing operations."""
    
    def __init__(self, config: Dict):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.ocr = OCRRecognizer(config)

    def extract_text(self, pdf_path: Path) -> str:
        """
        Extracts text content from PDF file.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            PDFProcessingError: If text extraction fails
        """
        try:
            if not pdf_path.exists():
                raise PDFProcessingError(f"PDF file not found: {pdf_path}")
            
            doc = PdfDocument()
            doc.LoadFromFile(str(pdf_path))
            
            text = []
            for page_num in range(min(2, doc.Pages.Count)):
                page = doc.Pages[page_num]
                text_extractor = PdfTextExtractor(page)
                extract_options = PdfTextExtractOptions()
                text.append(text_extractor.ExtractText(extract_options))
            
            return ' '.join(text)
        except Exception as e:
            raise PDFProcessingError(f"Failed to extract text: {str(e)}")

    def extract_images(self, pdf_path: Path) -> Tuple[str, str]:
        """
        Extracts images from PDF file.
        
        Returns:
            Tuple containing paths to registration number image and car photo
        """
        try:
            doc = PdfDocument()
            doc.LoadFromFile(str(pdf_path))
            
            page = doc.Pages[1]
            image_helper = PdfImageHelper()
            image_infos = image_helper.GetImagesInfo(page)
            
            save_dir = pdf_path.parent / "images"
            save_dir.mkdir(exist_ok=True)
            
            reg_filename = save_dir / f'number_{pdf_path.stem}.png'
            car_filename = save_dir / f'car_{pdf_path.stem}.png'
            
            image_infos[0].Image.Save(str(reg_filename))
            image_infos[1].Image.Save(str(car_filename))
            
            return str(reg_filename), str(car_filename)
        except Exception as e:
            raise PDFProcessingError(f"Failed to extract images: {str(e)}")

    def get_text_info(self, text: str) -> Dict:
        """
        Extracts relevant information from text content.
        
        Args:
            text: Extracted text content
            
        Returns:
            Dictionary containing extracted information
        """
        try:
            # Your existing text processing logic here
            # This is a placeholder - implement according to your specific needs
            info = {
                'ПОСТАНОВЛЕНИЕ': '',
                'дата_нарушения': '',
                'время_нарушения': '',
                'адрес': '',
                'номер_тс': '',
                'сумма_штрафа': '',
                'ИГР': ''
            }
            
            # Add your text parsing logic here
            
            return info
        except Exception as e:
            raise PDFProcessingError(f"Failed to extract information from text: {str(e)}")

    def tesseract_recognizer(self, image: np.ndarray) -> str:
        """
        Perform OCR on the preprocessed image.
        
        Args:
            image: Preprocessed image array
            
        Returns:
            str: Recognized text
        """
        try:
            return self.ocr.recognize(image)
        except Exception as e:
            self.logger.error(f"OCR recognition failed: {str(e)}")
            raise OCRError(f"Failed to recognize text: {str(e)}")

class ImageProcessor:
    """Handles image processing operations."""
    
    def __init__(self, config: Dict):
        self.config = config['image_processing']
        self.logger = logging.getLogger(__name__)

    def preprocess(self, image_path: str) -> np.ndarray:
        """
        Preprocesses image for OCR.
        
        Args:
            image_path: Path to image file
            
        Returns:
            Preprocessed image as numpy array
        """
        try:
            # Read image
            img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)
            if img is None:
                raise ImageProcessingError(f"Failed to load image: {image_path}")

            # Resize
            img = cv2.resize(
                img,
                None,
                fx=self.config['resize_factor'],
                fy=self.config['resize_factor'],
                interpolation=cv2.INTER_CUBIC
            )

            # Apply Gaussian blur
            img = cv2.GaussianBlur(
                img,
                (self.config['blur_kernel_size'], self.config['blur_kernel_size']),
                0
            )

            # Apply CLAHE
            clahe = cv2.createCLAHE(
                clipLimit=self.config['clahe']['clip_limit'],
                tileGridSize=tuple(self.config['clahe']['tile_grid_size'])
            )
            img = clahe.apply(img)

            # Denoise
            img = cv2.fastNlMeansDenoising(
                img,
                h=self.config['denoise']['h'],
                templateWindowSize=self.config['denoise']['template_window_size'],
                searchWindowSize=self.config['denoise']['search_window_size']
            )

            # Sharpen
            kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
            img = cv2.filter2D(img, -1, kernel)

            # Normalize
            img = cv2.normalize(img, None, 0, 255, cv2.NORM_MINMAX)

            # Threshold
            img = cv2.adaptiveThreshold(
                img,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY_INV,
                self.config['threshold']['block_size'],
                self.config['threshold']['C']
            )

            img = cv2.bitwise_not(img)
            img = cv2.normalize(img, None, 0, 255, cv2.NORM_MINMAX)
            img = cv2.medianBlur(img, 3)

            return img
        except Exception as e:
            raise ImageProcessingError(f"Image preprocessing failed: {str(e)}")

class DocumentGenerator:
    """Handles document generation operations."""
    
    def __init__(self, config: Dict):
        self.config = config['output']
        self.logger = logging.getLogger(__name__)

    def create_table_document(self, data: List[Dict], filename: Optional[str] = None) -> None:
        """Creates document with violation data table."""
        try:
            doc = Document()
            table = doc.add_table(rows=1, cols=len(data[0]))
            table.style = 'Table Grid'

            # Add headers
            for i, key in enumerate(data[0].keys()):
                table.cell(0, i).text = str(key)
                table.cell(0, i).paragraphs[0].runs[0].bold = True

            # Add data
            for item in data:
                row_cells = table.add_row().cells
                for i, (key, value) in enumerate(item.items()):
                    if key in ['car_photo_filename', 'regnumber_filename']:
                        try:
                            run = row_cells[i].paragraphs[0].add_run()
                            run.add_picture(value, width=Inches(self.config['docx']['image_width']))
                        except:
                            row_cells[i].text = "Image not available"
                    else:
                        row_cells[i].text = str(value)

            output_filename = filename or self.config['docx']['table_filename']
            doc.save(output_filename)
            self.logger.info(f"Table document created: {output_filename}")
            
        except Exception as e:
            raise ConfigurationError(f"Failed to create table document: {str(e)}")

    def create_daily_statistics(self, df: pd.DataFrame, filename: Optional[str] = None) -> None:
        """Creates daily statistics document."""
        try:
            stats_df = (df.groupby('адрес')
                         .size()
                         .reset_index(name='Количество нарушений')
                         .sort_values('Количество нарушений', ascending=False))
            
            stats_df.insert(0, '№', range(1, len(stats_df) + 1))
            
            doc = Document()
            doc.add_heading('Статистика нарушений по адресам', 0)
            doc.add_paragraph(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
            
            table = doc.add_table(rows=1, cols=len(stats_df.columns))
            table.style = 'Table Grid'
            
            # Add headers
            for i, column in enumerate(stats_df.columns):
                table.cell(0, i).text = str(column)
                table.cell(0, i).paragraphs[0].runs[0].bold = True
            
            # Add data
            for _, row in stats_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            
            output_filename = filename or self.config['docx']['stats_filename']
            doc.save(output_filename)
            self.logger.info(f"Statistics document created: {output_filename}")
            
        except Exception as e:
            raise ConfigurationError(f"Failed to create statistics document: {str(e)}")

class EmailSender:
    """Handles email operations."""
    
    def __init__(self, config: Dict):
        self.config = config['email']
        self.logger = logging.getLogger(__name__)

    def send_violation_alert(self, data: Dict, recipient_type: str = 'cargo') -> None:
        """Sends violation alert email."""
        try:
            msg = MIMEMultipart()
            msg['From'] = self.config['sender_email']
            msg['To'] = self.config['recipients'][recipient_type]
            msg['Subject'] = f"Уведомление о нарушении ПДД - {data['ПОСТАНОВЛЕНИЕ']}"
            
            body = self._create_violation_body(data)
            msg.attach(MIMEText(body, 'html'))
            
            if self.config['alerts']['include_images']:
                self._attach_images(msg, data)
            
            self._send_email(msg)
            self.logger.info(f"Violation alert sent to {recipient_type}")
            
        except Exception as e:
            raise EmailError(f"Failed to send violation alert: {str(e)}")

    def send_daily_statistics(self, df: pd.DataFrame) -> None:
        """Sends daily statistics email."""
        try:
            msg = MIMEMultipart()
            msg['From'] = self.config['sender_email']
            msg['To'] = self.config['recipients']['cargo']
            msg['Subject'] = f"Дневная статистика нарушений ПДД - {datetime.now().strftime('%d.%m.%Y')}"
            
            stats_df = (df.groupby('адрес')
                         .size()
                         .reset_index(name='Количество нарушений')
                         .sort_values('Количество нарушений', ascending=False))
            
            body = self._create_statistics_body(stats_df)
            msg.attach(MIMEText(body, 'html'))
            
            self._send_email(msg)
            self.logger.info("Daily statistics email sent")
            
        except Exception as e:
            raise EmailError(f"Failed to send daily statistics: {str(e)}")

    def _create_violation_body(self, data: Dict) -> str:
        """Creates HTML body for violation email."""
        return f"""
        <html>
            <body>
                <h2>Обнаружено нарушение ПДД</h2>
                <table style="border-collapse: collapse; width: 100%;">
                    <tr>
                        <th style="border: 1px solid black; padding: 8px;">Параметр</th>
                        <th style="border: 1px solid black; padding: 8px;">Значение</th>
                    </tr>
                    <tr>
                        <td style="border: 1px solid black; padding: 8px;">Постановление</td>
                        <td style="border: 1px solid black; padding: 8px;">{data['ПОСТАНОВЛЕНИЕ']}</td>
                    </tr>
                    <tr>
                        <td style="border: 1px solid black; padding: 8px;">Дата нарушения</td>
                        <td style="border: 1px solid black; padding: 8px;">{data['дата_нарушения']}</td>
                    </tr>
                    <tr>
                        <td style="border: 1px solid black; padding: 8px;">Время нарушения</td>
                        <td style="border: 1px solid black; padding: 8px;">{data['время_нарушения']}</td>
                    </tr>
                    <tr>
                        <td style="border: 1px solid black; padding: 8px;">Адрес</td>
                        <td style="border: 1px solid black; padding: 8px;">{data['адрес']}</td>
                    </tr>
                    <tr>
                        <td style="border: 1px solid black; padding: 8px;">Номер ТС</td>
                        <td style="border: 1px solid black; padding: 8px;">{data['номер_тс']}</td>
                    </tr>
                    <tr>
                        <td style="border: 1px solid black; padding: 8px;">Сумма штрафа</td>
                        <td style="border: 1px solid black; padding: 8px;">{data['сумма_штрафа']} руб.</td>
                    </tr>
                </table>
            </body>
        </html>
        """

    def _create_statistics_body(self, stats_df: pd.DataFrame) -> str:
        """Creates HTML body for statistics email."""
        total = stats_df['Количество нарушений'].sum()
        
        table_rows = ""
        for _, row in stats_df.iterrows():
            percentage = (row['Количество нарушений'] / total * 100).round(2)
            table_rows += f"""
                <tr>
                    <td style="border: 1px solid black; padding: 8px;">{row['адрес']}</td>
                    <td style="border: 1px solid black; padding: 8px;">{row['Количество нарушений']}</td>
                    <td style="border: 1px solid black; padding: 8px;">{percentage}%</td>
                </tr>
            """
            
        return f"""
        <html>
            <body>
                <h2>Дневная статистика нарушений ПДД</h2>
                <p>Дата: {datetime.now().strftime('%d.%m.%Y')}</p>
                <p>Общее количество нарушений: {total}</p>
                <table style="border-collapse: collapse; width: 100%;">
                    <tr>
                        <th style="border: 1px solid black; padding: 8px;">Адрес</th>
                        <th style="border: 1px solid black; padding: 8px;">Количество</th>
                        <th style="border: 1px solid black; padding: 8px;">Процент</th>
                    </tr>
                    {table_rows}
                </table>
            </body>
        </html>
        """

    def _attach_images(self, msg: MIMEMultipart, data: Dict) -> None:
        """Attaches images to email message."""
        for key in ['car_photo_filename', 'regnumber_filename']:
            if key in data and data[key]:
                try:
                    with open(data[key], 'rb') as f:
                        img = MIMEImage(f.read())
                        img.add_header('Content-ID', f'<{key}>')
                        msg.attach(img)
                except Exception as e:
                    self.logger.warning(f"Could not attach image {data[key]}: {str(e)}")

    def _send_email(self, msg: MIMEMultipart) -> None:
        """Sends email using SMTP."""
        try:
            with smtplib.SMTP(self.config['smtp_server'], self.config['smtp_port']) as server:
                server.starttls()
                server.login(
                    self.config['sender_email'],
                    self.config['sender_password']
                )
                server.send_message(msg)
        except Exception as e:
            raise EmailError(f"Failed to send email: {str(e)}")

